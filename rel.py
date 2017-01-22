from docx import Document

document = Document('merged.docx')
tables = document.tables
for table in tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                print(paragraph.text)
#for p in document.paragraphs:
#    print (p.text)
