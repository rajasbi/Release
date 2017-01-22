import sys
from docx import Document

document = Document('one.docx')
tables = document.tables
for row in tables[1].rows:
    print row.cells[0].paragraphs[0].text + row.cells[1].paragraphs[0].text
